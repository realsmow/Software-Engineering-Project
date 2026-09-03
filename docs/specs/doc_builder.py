"""Turn a content module (srs_data.py, sds_data.py) into a .docx and a .pdf.

Used by gen_srs_docx.py and gen_sds_docx.py — see specs/README.md for how to run.

Three things here are less obvious than they look:

1. Thai text is "complex script" to Word, which takes its font and size from the
   *cs* attributes (w:cs, w:szCs, w:bCs) and ignores the Latin ones. Every helper
   sets both, or Word renders the Thai runs in a substituted font at the wrong
   size while the English runs look fine.

2. Column widths have to be written to the table grid (w:gridCol), not only to
   the cells: with a fixed layout, Word and LibreOffice lay out from the grid and
   quietly ignore per-cell widths.

3. The table of contents is built twice. Pass one produces a PDF with the right
   number of TOC lines but dummy page numbers; we read back which page each
   heading landed on, then rebuild with the real numbers. Word still gets a live
   TOC field (it refreshes on open), while the PDF — which LibreOffice renders
   from the field's cached result, never recalculating it — comes out correct.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import (WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT,
                            WD_TAB_LEADER)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# --- look and feel (shared by every document in this folder) -----------------

FONT = "Sarabun"          # the product's own UI font, and what the source docs used

SIZE_BODY = 14
SIZE_TABLE = 13
SIZE_H1 = 18
SIZE_H2 = 15
SIZE_SMALL = 11

HEADER_FILL = "E8E8E8"    # grey, not colour: these are printed in black and white
RULE_COLOR = "999999"
MUTED = "555555"

PAGE_W, PAGE_H = 21.0, 29.7          # A4 in cm
MARGIN = 2.5
CONTENT_W = PAGE_W - 2 * MARGIN      # 16 cm — the TOC right tab stop

OUT_DIR = Path(__file__).resolve().parent.parent / "build"


# --- low-level OOXML helpers -------------------------------------------------

def _el(tag: str, **attrs) -> OxmlElement:
    """Create a w:-namespaced element with w:-namespaced attributes."""
    node = OxmlElement(f"w:{tag}")
    for key, value in attrs.items():
        node.set(qn(f"w:{key}"), value)
    return node


def _apply_font(rpr, size: int, *, bold: bool, color: str | None, name: str) -> None:
    rpr.append(_el("rFonts", ascii=name, hAnsi=name, cs=name, eastAsia=name))
    if bold:
        rpr.append(_el("b", val="1"))
        rpr.append(_el("bCs", val="1"))
    rpr.append(_el("sz", val=str(size * 2)))       # OOXML counts half-points
    rpr.append(_el("szCs", val=str(size * 2)))
    if color:
        rpr.append(_el("color", val=color))


def style_font(rpr, size: int, *, bold: bool = False, color: str | None = None,
               name: str = FONT) -> None:
    """Font, size and weight for a style — Latin and complex-script alike."""
    _apply_font(rpr, size, bold=bold, color=color, name=name)
    # tell Word the complex-script language is Thai so it breaks lines correctly
    rpr.append(_el("lang", val="en-US", bidi="th-TH"))


def run_font(run, size: int, *, bold: bool = False, color: str | None = None,
             name: str = FONT) -> None:
    """Same, for a single run (used where no paragraph style fits)."""
    _apply_font(run._r.get_or_add_rPr(), size, bold=bold, color=color, name=name)


def add_field(paragraph, instruction: str, shown_text: str, size: int) -> list:
    """Open a Word field. Returns the runs so a caller can close it later.

    A field is begin / instruction / separate / cached result / end. Splitting it
    up matters for the TOC, whose cached result spans many paragraphs.
    """
    begin = paragraph.add_run()
    begin._r.append(_el("fldChar", fldCharType="begin"))

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run._r.append(instr)

    separate = paragraph.add_run()
    separate._r.append(_el("fldChar", fldCharType="separate"))

    runs = [begin, instr_run, separate]
    if shown_text:
        shown = paragraph.add_run(shown_text)
        run_font(shown, size)
        runs.append(shown)
    for run in runs:
        run_font(run, size)
    return runs


def close_field(paragraph, size: int) -> None:
    end = paragraph.add_run()
    end._r.append(_el("fldChar", fldCharType="end"))
    run_font(end, size)


def simple_field(paragraph, instruction: str, shown_text: str, size: int) -> None:
    add_field(paragraph, instruction, shown_text, size)
    close_field(paragraph, size)


def enable_field_update(document) -> None:
    """Ask the reader to recalculate fields on open, so Word refreshes the TOC.

    settings.xml is a strict sequence, so this has to be inserted before the
    elements that follow it in the schema rather than simply appended.
    """
    settings = document.settings.element
    later = {qn(f"w:{t}") for t in
             ("hdrShapeDefaults", "footnotePr", "endnotePr", "compat", "docVars",
              "rsids", "themeFontLang", "clrSchemeMapping", "shapeDefaults",
              "decimalSymbol", "listSeparator")}
    node = _el("updateFields", val="true")
    for child in settings:
        if child.tag in later:
            child.addprevious(node)
            return
    settings.append(node)


def bottom_rule(paragraph, color: str = RULE_COLOR) -> None:
    borders = OxmlElement("w:pBdr")
    borders.append(_el("bottom", val="single", sz="6", space="4", color=color))
    paragraph._p.get_or_add_pPr().append(borders)


# --- document skeleton -------------------------------------------------------

def setup_styles(document) -> None:
    """Rewrite the built-in styles instead of formatting run by run.

    Heading 1/2 stay the built-in styles on purpose: a TOC field collects
    paragraphs by style, so custom heading styles would come out empty in Word.
    """
    normal = document.styles["Normal"]
    style_font(normal.element.get_or_add_rPr(), SIZE_BODY)
    fmt = normal.paragraph_format
    fmt.space_after = Pt(8)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.15

    for name, size, space_before in (("Heading 1", SIZE_H1, 20),
                                     ("Heading 2", SIZE_H2, 14)):
        style = document.styles[name]
        rpr = style.element.get_or_add_rPr()
        # the built-in headings arrive blue and semi-italic; clear that first
        for tag in ("rFonts", "b", "bCs", "sz", "szCs", "color", "i", "iCs", "lang"):
            for old in rpr.findall(qn(f"w:{tag}")):
                rpr.remove(old)
        style_font(rpr, size, bold=True, color="000000")
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def setup_page(document, running_header: str) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Cm(PAGE_W), Cm(PAGE_H)
    section.left_margin = section.right_margin = Cm(MARGIN)
    section.top_margin = section.bottom_margin = Cm(MARGIN)
    section.header_distance = section.footer_distance = Cm(1.25)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_font(header.add_run(running_header), SIZE_SMALL, color=MUTED)
    bottom_rule(header)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_font(footer.add_run("หน้า "), SIZE_SMALL, color=MUTED)
    simple_field(footer, " PAGE ", "1", SIZE_SMALL)
    run_font(footer.add_run(" / "), SIZE_SMALL, color=MUTED)
    simple_field(footer, " NUMPAGES ", "1", SIZE_SMALL)


def add_cover(document, cover: dict) -> None:
    def centered(text: str, size: int, *, bold: bool = False,
                 before: int = 0, after: int = 6, color: str | None = None):
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run_font(paragraph.add_run(text), size, bold=bold, color=color)

    centered("", SIZE_BODY, before=96)
    centered(cover["kicker"], 20, bold=True, after=2)
    centered(cover["acronym"], 18, bold=True, after=30, color=MUTED)
    centered(cover["title_th"], 22, bold=True, after=6)
    centered(cover["title_en"], 15, after=64)
    centered(cover["version"], 16, bold=True, after=4)
    centered(cover["date"], 16, after=58)
    centered(cover["team"], SIZE_BODY, after=3)
    centered(cover["course"], SIZE_BODY, after=3)
    centered(cover["faculty"], SIZE_BODY, after=3)
    centered(cover["university"], SIZE_BODY, after=0)

    document.add_page_break()


def add_toc(document, title: str, headings: list[tuple[int, str]],
            pages: dict[str, int] | None) -> None:
    """A real TOC field whose cached result we fill in ourselves.

    `pages` is None on the first pass — the entries are then written with a dummy
    number, which keeps the line count (and therefore the pagination) identical
    to the final document.
    """
    heading = document.add_paragraph()
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.keep_with_next = True
    # deliberately not a Heading style: the contents should not list itself
    run_font(heading.add_run(title), SIZE_H1, bold=True)

    first = document.add_paragraph()
    add_field(first, r'TOC \o "1-2" \h \z \u', "", SIZE_BODY)

    paragraph = first
    for index, (level, text) in enumerate(headings):
        if index:                       # the field opened on `first`
            paragraph = document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.space_after = Pt(2)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        fmt.left_indent = Cm(0.0 if level == 1 else 0.8)
        fmt.tab_stops.add_tab_stop(Cm(CONTENT_W), WD_TAB_ALIGNMENT.RIGHT,
                                   WD_TAB_LEADER.DOTS)
        page_no = "0" if pages is None else str(pages.get(text, 0))
        run_font(paragraph.add_run(f"{text}\t{page_no}"), SIZE_BODY, bold=(level == 1))

    close_field(paragraph, SIZE_BODY)
    document.add_page_break()


# --- content blocks ----------------------------------------------------------

def add_table(document, widths_cm, header, rows) -> None:
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    table.autofit = False
    table._tbl.tblPr.append(_el("tblLayout", type="fixed"))

    # the grid is what a fixed layout actually measures from; cell widths alone
    # are ignored by both Word and LibreOffice
    for column, width in zip(table.columns, widths_cm):
        column.width = Cm(width)

    def fill(cell, text: str, *, bold: bool, width_cm: float):
        cell.width = Cm(width_cm)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run_font(paragraph.add_run(text), SIZE_TABLE, bold=bold)

    header_row = table.rows[0]
    # repeat the header on every page the table spills onto
    header_row._tr.get_or_add_trPr().append(_el("tblHeader", val="true"))
    for cell, text, width in zip(header_row.cells, header, widths_cm):
        fill(cell, text, bold=True, width_cm=width)
        cell._tc.get_or_add_tcPr().append(
            _el("shd", val="clear", color="auto", fill=HEADER_FILL))

    for values in rows:
        row = table.add_row()
        # keep a requirement and its priority together on one page
        row._tr.get_or_add_trPr().append(_el("cantSplit", val="true"))
        for cell, text, width in zip(row.cells, values, widths_cm):
            fill(cell, text, bold=False, width_cm=width)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    spacer.paragraph_format.line_spacing = Pt(6)


def add_bullets(document, items) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph.paragraph_format.line_spacing = 1.15
        run_font(paragraph.add_run(item), SIZE_BODY)


def add_code(document, lines) -> None:
    """A monospaced block with a light background — folder trees, snippets."""
    for index, line in enumerate(lines):
        paragraph = document.add_paragraph()
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(6 if index == 0 else 0)
        fmt.space_after = Pt(6 if index == len(lines) - 1 else 0)
        fmt.left_indent = Cm(0.4)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph._p.get_or_add_pPr().append(
            _el("shd", val="clear", color="auto", fill="F2F2F2"))
        run_font(paragraph.add_run(line or " "), 10, name="DejaVu Sans Mono")


def add_image(document, rel_path: str, width_cm: float, caption: str) -> None:
    """A centred figure with a caption underneath.

    Paths are relative to this folder so the content files stay portable.
    Width is given in cm rather than scaled to the page, because a wide BPMN
    strip and a tall use-case diagram need very different treatment.
    """
    image = Path(__file__).resolve().parent / rel_path
    if not image.exists():
        raise FileNotFoundError(f"figure not found: {image}")

    holder = document.add_paragraph()
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    holder.paragraph_format.space_before = Pt(8)
    holder.paragraph_format.space_after = Pt(3)
    holder.paragraph_format.keep_with_next = True
    holder.add_run().add_picture(str(image), width=Cm(width_cm))

    if caption:
        label = document.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label.paragraph_format.space_after = Pt(12)
        run_font(label.add_run(caption), SIZE_TABLE, color=MUTED)


def _new_section(document, *, landscape: bool):
    """Start a page-break section, portrait or landscape.

    python-docx copies the previous section's page setup, so the width and
    height have to be swapped by hand — setting `orientation` alone changes
    only the flag Word reports, not the printed page.
    """
    section = document.add_section(WD_SECTION.NEW_PAGE)
    long_side, short_side = Cm(PAGE_H), Cm(PAGE_W)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = long_side, short_side
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width, section.page_height = short_side, long_side
    section.left_margin = section.right_margin = Cm(MARGIN)
    section.top_margin = section.bottom_margin = Cm(MARGIN)
    section.header_distance = section.footer_distance = Cm(1.25)
    return section


def add_wide_image(document, rel_path: str, caption: str) -> None:
    """A diagram too wide to read on a portrait page: give it a landscape page.

    A BPMN strip is roughly 2.4 : 1, so on a portrait page its labels shrink to
    the point of being decorative. Turning the page sideways buys ~60% more
    width, which is the difference between a reader following the flow and just
    seeing that a flow exists.
    """
    _new_section(document, landscape=True)
    add_image(document, rel_path, PAGE_H - 2 * MARGIN, caption)
    _new_section(document, landscape=False)


def add_body(document, body) -> None:
    for block in body:
        kind = block[0]
        if kind == "h1":
            document.add_paragraph(block[1], style="Heading 1")
        elif kind == "h2":
            document.add_paragraph(block[1], style="Heading 2")
        elif kind == "p":
            document.add_paragraph(block[1])
        elif kind == "note":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.space_after = Pt(10)
            run_font(paragraph.add_run(block[1]), SIZE_TABLE, color="404040")
        elif kind == "table":
            add_table(document, block[1], block[2], block[3])
        elif kind == "bullets":
            add_bullets(document, block[1])
        elif kind == "code":
            add_code(document, block[1])
        elif kind == "image":
            add_image(document, block[1], block[2], block[3])
        elif kind == "wide_image":
            add_wide_image(document, block[1], block[2])
        else:
            raise ValueError(f"unknown block type: {kind!r}")


def set_properties(document, data) -> None:
    props = document.core_properties
    props.title = data.DOC_TITLE
    props.subject = "University Equipment Lending Management System (ULMs)"
    props.author = data.COVER["team"]
    props.category = data.COVER["kicker"]
    props.comments = f"สร้างจาก docs/specs/{data.GENERATOR} — แก้เนื้อหาที่ {data.SOURCE_FILE}"
    props.language = "th-TH"


# --- assembly and export -----------------------------------------------------

def headings_of(body) -> list[tuple[int, str]]:
    """The h1/h2 blocks, in order — the entries a TOC field would collect."""
    return [(1 if block[0] == "h1" else 2, block[1])
            for block in body if block[0] in ("h1", "h2")]


def render(data, pages: dict[str, int] | None) -> Document:
    document = Document()
    setup_styles(document)
    setup_page(document, data.RUNNING_HEADER)
    add_cover(document, data.COVER)
    add_toc(document, data.TOC_TITLE, headings_of(data.BODY), pages)
    add_body(document, data.BODY)
    set_properties(document, data)
    enable_field_update(document)
    return document


def to_pdf(docx_path: Path, out_dir: Path) -> Path:
    """Convert with LibreOffice in a throwaway profile, so a running Writer is safe."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) is not installed — cannot make a PDF")

    with tempfile.TemporaryDirectory(prefix="specs-soffice-") as profile:
        result = subprocess.run(
            [soffice, f"-env:UserInstallation=file://{profile}",
             "--headless", "--norestore", "--convert-to", "pdf:writer_pdf_Export",
             "--outdir", str(out_dir), str(docx_path)],
            capture_output=True, text=True, timeout=600,
        )
    if result.returncode != 0:
        print(result.stdout, result.stderr, file=sys.stderr)
        raise RuntimeError("LibreOffice failed to convert the document")

    pdf = out_dir / (docx_path.stem + ".pdf")
    if not pdf.exists():
        raise RuntimeError(f"expected {pdf} but LibreOffice produced nothing")
    return pdf


def _page_texts(pdf: Path) -> list[str]:
    pages = subprocess.run(["pdftotext", str(pdf), "-"],
                           capture_output=True, text=True, check=True).stdout
    return pages.split("\f")


def resolve_pages(pdf: Path, headings, first_body_text: str) -> dict[str, int]:
    """Read back which printed page each heading landed on.

    Thai has no word spaces and the extractor wraps lines, so both sides are
    stripped of whitespace before matching. The scan starts at the first page of
    real content — otherwise every heading would match on the contents page — and
    only moves forward, so a heading can never resolve to an earlier page than
    the one before it.
    """
    squash = lambda text: re.sub(r"\s+", "", text)
    pages = [squash(text) for text in _page_texts(pdf)]

    needle = squash(first_body_text)[:40]
    body_start = next((i for i, text in enumerate(pages) if needle in text), 0)

    found: dict[str, int] = {}
    cursor = body_start
    for _, text in headings:
        target = squash(text)
        for index in range(cursor, len(pages)):
            if target in pages[index]:
                found[text] = index + 1        # pdftotext pages are 0-based here
                cursor = index
                break
    return found


def build(data) -> tuple[Path, Path]:
    """Two passes: measure the page numbers, then write the real document."""
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    check_font()

    with tempfile.TemporaryDirectory(prefix="specs-pass1-") as scratch:
        scratch_dir = Path(scratch)
        draft = scratch_dir / f"{data.OUTPUT_STEM}.docx"
        render(data, pages=None).save(draft)
        draft_pdf = to_pdf(draft, scratch_dir)
        first_body_text = next(block[1] for block in data.BODY if block[0] == "p")
        pages = resolve_pages(draft_pdf, headings_of(data.BODY), first_body_text)

    missing = [text for _, text in headings_of(data.BODY) if text not in pages]
    if missing:
        print(f"! no page number found for {len(missing)} heading(s): {missing[:3]}",
              file=sys.stderr)

    docx_path = OUT_DIR / f"{data.OUTPUT_STEM}.docx"
    render(data, pages=pages).save(docx_path)
    pdf_path = to_pdf(docx_path, OUT_DIR)

    print(f"wrote {docx_path}  ({docx_path.stat().st_size // 1024} KB)")
    print(f"wrote {pdf_path}  ({pdf_path.stat().st_size // 1024} KB)")
    return docx_path, pdf_path


def check_font() -> None:
    fc_list = shutil.which("fc-list")
    if not fc_list:
        return
    families = subprocess.run([fc_list, ":", "family"], capture_output=True, text=True)
    if FONT.lower() not in families.stdout.lower():
        print(f"! {FONT} is not installed — Word and LibreOffice will substitute "
              f"another font and the layout will shift.\n"
              f"  Install it from https://fonts.google.com/specimen/{FONT}",
              file=sys.stderr)
